import os
import random
import nltk
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from tavily import TavilyClient
from docx import Document
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from nltk.corpus import wordnet
from rake_nltk import Rake
from prompts import blog_prompt, summary_prompt, caption_prompt

# Load environment variables from .env file
load_dotenv()

# Download NLTK data
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')

class ArticleProcessor:
    def __init__(self):
        # Get API keys from environment variables
        tavily_api_key = os.getenv("TAVILY_API_KEY")
        openai_api_key = os.getenv("OPENAI_API_KEY")

        if not tavily_api_key or not openai_api_key:
            raise ValueError("API keys must be provided in the .env file")

        # Set up API keys and initialize the LLM
        os.environ["TAVILY_API_KEY"] = tavily_api_key
        os.environ["OPENAI_API_KEY"] = openai_api_key
        self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.7)

        # Define prompt templates
        self.prompt_templates = {
            "summary": PromptTemplate(
                input_variables=["text"],
                template=summary_prompt
            ),
            "blog": PromptTemplate(
                input_variables=["summary", "text"],
                template=blog_prompt
            ),
            "caption": PromptTemplate(
                input_variables=["summary"],
                template=caption_prompt
            )
        }

        # Initialize TavilyClient
        self.client = TavilyClient(api_key=tavily_api_key)

        # Initialize tokenizer and text splitter
        self.tokenizer = tiktoken.encoding_for_model("gpt-4o-mini")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=8000,  # Reduced chunk size to ensure smaller chunks
            chunk_overlap=200,
            length_function=lambda text: len(self.tokenizer.encode(text))
        )

    def search_articles(self, query, max_results=5):
        # Execute search query with TavilyClient
        response = self.client.search(query, max_results=max_results, include_raw_content=True)

        # Prepare results
        articles = [(result['url'], result['raw_content']) for result in response['results']]
        content = "\n".join(article[1] for article in articles)
        urls = [article[0] for article in articles]

        return articles, content, urls

    def create_chain(self, prompt_template):
        return LLMChain(llm=self.llm, prompt=prompt_template)

    def process_text(self, text, prompt_type):
        chain = self.create_chain(self.prompt_templates[prompt_type])
        return chain.run(text)

    def save_text_to_word(self, text, filename):
        doc = Document()
        doc.add_heading('Blog Post', 0)
        doc.add_paragraph(text)
        doc.save(filename)

    def paraphrase_sentence(self, sentence, excluded_words=[], topic_words=[], vocabulary=[]):
        words = nltk.word_tokenize(sentence)
        paraphrased_words = []

        for word in words:
            synonyms = wordnet.synsets(word)
            if synonyms:
                synonyms = [synonym.lemmas()[0].name() for synonym in synonyms if synonym.lemmas()[0].name() not in excluded_words and synonym.lemmas()[0].name() in vocabulary]
                if topic_words:
                    synonyms = [synonym for synonym in synonyms if any(topic_word in wordnet.synsets(synonym)[0].definition() for topic_word in topic_words)]
                if synonyms:
                    synonym = random.choice(synonyms)
                    paraphrased_words.append(synonym)
                else:
                    paraphrased_words.append(word)
            else:
                paraphrased_words.append(word)

        paraphrased_sentence = ' '.join(paraphrased_words)
        return paraphrased_sentence

    def humanize_text(self, text, excluded_words=[], topic_words=[], vocabulary=[]):
        # Split text based on "#" or "*"
        paragraphs = text.split('\n\n')
        humanized_paragraphs = []

        for paragraph in paragraphs:
            # Check if the paragraph should be split
            if paragraph.startswith("#") or paragraph.startswith("*"):
                humanized_paragraphs.append(paragraph)
            else:
                sentences = nltk.sent_tokenize(paragraph)
                humanized_sentences = []

                for sentence in sentences:
                    humanized_sentence = self.paraphrase_sentence(sentence, excluded_words, topic_words, vocabulary)
                    humanized_sentences.append(humanized_sentence)

                humanized_paragraphs.append(' '.join(humanized_sentences))

        humanized_text = '\n\n'.join(humanized_paragraphs)
        return humanized_text

    def process_query(self, query):
        # Get articles content
        _, content, _ = self.search_articles(query)

        # Wrap content in Document objects
        documents = [LangChainDocument(page_content=content)]

        # Split the content if it's too long
        chunks = self.text_splitter.split_documents(documents)
        summarized_chunks = []

        for chunk in chunks:
            chunk_content = chunk.page_content  # Access the page_content attribute
            summary = self.process_text(chunk_content, "summary")
            summarized_chunks.append(summary)

        # Combine all summaries
        full_summary = "\n".join(summarized_chunks)

        # If the full summary is still too long, summarize it further
        while len(self.tokenizer.encode(full_summary)) > 10000:  # Adjust this threshold based on your needs
            full_summary = self.process_text(full_summary, "summary")

        # Generate a new blog based on the summary
        blog_post = self.process_text({"summary": full_summary, "text": content}, "blog")
        # Generate a caption based on the summary
        caption = self.process_text(summary, "caption")
        print("\nGenerated Caption:")
        print(caption)
        
        # Paraphrase the generated blog post
        r = Rake()
        r.extract_keywords_from_text(blog_post)
        keywords = r.get_ranked_phrases()

        vocabulary = ["AI", "artificial intelligence", "software development", "coding", "developer", "technology", "innovation", "efficiency", "learning", "skills", "certificate", "opportunities", "benefits", "job market", "cutting-edge", "tools", "platforms", "code", "user", "needs", "requirements", "design", "testing", "reliability", "black box", "biases", "human creativity", "computational prowess", "productivity", "effectiveness", "projects", "learning platform", "real-world scenarios", "prompt engineering", "AI-generated code", "job displacement", "ethical dilemmas", "future", "possibilities", "synergy", "trend", "landscape", "embrace", "transformation", "boost", "flex", "skills", "stand out", "LinkedIn", "certificate", "shiny", "opportunities", "endeavor", "journey", "coding", "brave new world", "game-changer", "edge"]
        topic_words = ["AI", "software development", "learning", "skills", "opportunities", "benefits", "job market", "tools", "platforms", "code", "user", "needs", "requirements", "design", "testing", "reliability", "biases", "human creativity", "computational prowess", "productivity", "effectiveness", "projects", "learning platform", "real-world scenarios", "prompt engineering", "AI-generated code", "job displacement", "ethical dilemmas", "future", "possibilities", "synergy", "trend", "landscape", "embrace", "transformation", "boost", "flex", "skills", "stand out", "LinkedIn", "certificate", "shiny", "opportunities", "endeavor", "journey", "coding", "brave new world", "game-changer", "edge"]

        excluded_words = ["arsenic", "angstrom"]
        humanized_blog_post = self.humanize_text(blog_post, excluded_words, topic_words, vocabulary)

        # Save the paraphrased blog post to a Word document after paraphrasing
        filename_after = f"{query}_after.docx"
        self.save_text_to_word(humanized_blog_post, filename_after)

        print(f"Paraphrased blog post saved as {filename_after}")

if __name__ == "__main__":
    processor = ArticleProcessor()

    # User input for the query
    query = input("Enter the topic name: ")
    processor.process_query(query)
