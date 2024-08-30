import os
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from tavily import TavilyClient
from docx import Document


class CaptionProcessor:
    def __init__(self, tavily_api_key, openai_api_key):
        # Set up API keys
        os.environ["TAVILY_API_KEY"] = tavily_api_key
        os.environ["OPENAI_API_KEY"] = openai_api_key

        # Initialize the LLM
        self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.7)

        # Define the prompt templates
        self.prompt_templates = {
            "summary": PromptTemplate(
                input_variables=["text"],
                template="Find the key points and updated info in the following text:\n"
                         "1. The summary should be precise and paraphrased.\n"
                         "2. The system should generate an SEO-optimized summary.\n"
                         "Text:\n\n{text}\n\nSummary:"
            ),
            "caption": PromptTemplate(
                input_variables=["summary"],
                template="Create a caption based on the following summary:\n"
                         "1. The caption should be of 3 to 4 lines and must be SEO-optimized.\n"
                         "2. Use relevant hashtags that should also be SEO-optimized and trending.\n"
                         "Summary:\n\n{summary}\n\nCaption:"
            )
        }

        # Initialize TavilyClient
        self.client = TavilyClient(api_key=tavily_api_key)

    def get_articles_content(self, query, max_results=5):
        # Execute search query with TavilyClient
        response = self.client.search(query, max_results=max_results, include_raw_content=True)

        # Prepare results
        articles = []
        content = ""
        urls = []

        for result in response['results']:
            url = result['url']
            article_content = result['raw_content']
            articles.append([url, article_content])
            content += article_content + "\n"
            urls.append(url)

        return articles, content, urls

    def create_chain(self, prompt_template):
        return LLMChain(llm=self.llm, prompt=prompt_template)

    def process_text(self, text, prompt_type):
        chain = self.create_chain(self.prompt_templates[prompt_type])
        return chain.run(text)

    def save_text_to_word(self, text, filename, heading):
        doc = Document()
        doc.add_heading(heading, 0)
        doc.add_paragraph(text)
        doc.save(filename)

    def process_query(self, query):
        # Get articles content
        _, content, _ = self.get_articles_content(query)

        # Summarize the text
        summary = self.process_text(content, "summary")
        print("\nSummary:")
        print(summary)

        # Generate a caption based on the summary
        caption = self.process_text(summary, "caption")
        print("\nGenerated Caption:")
        print(caption)

        # Save caption to Word document
        caption_filename = f"{query}_caption.docx"
        self.save_text_to_word(caption, caption_filename, "Caption")
        print(f"\nCaption saved as {caption_filename}")

        # Save summary to Word document
        summary_filename = f"summary_{query}.docx"
        self.save_text_to_word(summary, summary_filename, "Summary")
        print(f"\nSummary saved as {summary_filename}")

if __name__ == "__main__":
    # Initialize CaptionProcessor with API keys
    tavily_api_key = "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"
    openai_api_key = "Your_openAI_key"

    processor = CaptionProcessor(tavily_api_key, openai_api_key)

    # User input for the query
    query = input("Enter the topic name: ")
    processor.process_query(query)
