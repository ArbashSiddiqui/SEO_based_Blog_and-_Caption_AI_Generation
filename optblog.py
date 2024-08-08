import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from tavily import TavilyClient
from docx import Document
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from prompts import blog_prompt, summary_prompt

# Load environment variables from .env file
load_dotenv()

class ArticleProcessor:
    def __init__(self):
        # Get API keys from environment variables
        tavily_api_key = os.getenv("TAVILY_API_KEY")
        openai_api_key = os.getenv("OPENAI_API_KEY")

        if not tavily_api_key or not openai_api_key:
            raise ValueError("API keys must be provided in the .env file")

        # Set up API keys and initialize the LLM
        os.environ["TAVILY_API_KEY"] = tavily_api_key
        os.environ["OPENAI_API_KEY"] = openai_api_key
        self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.7)

        # Define prompt templates
        self.prompt_templates = {
            "summary": PromptTemplate(
                input_variables=["text"],
                template=summary_prompt
            ),
            "blog": PromptTemplate(
                input_variables=["summary", "text"],
                template=blog_prompt
            )
        }

        # Initialize TavilyClient
        self.client = TavilyClient(api_key=tavily_api_key)

        # Initialize tokenizer and text splitter
        self.tokenizer = tiktoken.encoding_for_model("gpt-4o-mini")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=8000,  # Reduced chunk size to ensure smaller chunks
            chunk_overlap=200,
            length_function=lambda text: len(self.tokenizer.encode(text))
        )

    def search_articles(self, query, max_results=5):
        # Execute search query with TavilyClient
        response = self.client.search(query, max_results=max_results, include_raw_content=True)

        # Prepare results
        articles = [(result['url'], result['raw_content']) for result in response['results']]
        content = "\n".join(article[1] for article in articles)
        urls = [article[0] for article in articles]

        return articles, content, urls

    def create_chain(self, prompt_template):
        return LLMChain(llm=self.llm, prompt=prompt_template)

    def process_text(self, text, prompt_type):
        chain = self.create_chain(self.prompt_templates[prompt_type])
        return chain.run(text)

    def save_text_to_word(self, text, filename):
        doc = Document()
        doc.add_heading('Blog Post', 0)
        doc.add_paragraph(text)
        doc.save(filename)

    def process_query(self, query):
        # Get articles content
        _, content, _ = self.search_articles(query)

        # Wrap content in Document objects
        documents = [LangChainDocument(page_content=content)]

        # Split the content if it's too long
        chunks = self.text_splitter.split_documents(documents)
        summarized_chunks = []

        for chunk in chunks:
            chunk_content = chunk.page_content  # Access the page_content attribute
            summary = self.process_text(chunk_content, "summary")
            summarized_chunks.append(summary)

        # Combine all summaries
        full_summary = "\n".join(summarized_chunks)

        # If the full summary is still too long, summarize it further
        while len(self.tokenizer.encode(full_summary)) > 10000:  # Adjust this threshold based on your needs
            full_summary = self.process_text(full_summary, "summary")

        # Generate a new blog based on the summary
        blog_post = self.process_text({"summary": full_summary, "text": content}, "blog")
        print("\nGenerated Blog Post:")
        print(blog_post)

        # Save to Word document with the query as filename
        filename = f"{query}.docx"
        self.save_text_to_word(blog_post, filename)
        print(f"\nBlog post saved as {filename}")

if __name__ == "__main__":
    processor = ArticleProcessor()

    # User input for the query
    query = input("Enter the topic name: ")
    processor.process_query(query)
